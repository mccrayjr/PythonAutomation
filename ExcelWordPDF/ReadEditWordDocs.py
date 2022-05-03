#! /usr/bin/env python3
import docx

d =docx.Document('path/to/doc') #create a new file
d.paragraphs # returns an array of paragraphs - every time "enter" is hit
p = d.paragraphs[1]
p.runs #each paragraphs has "sentences" or runs which also change based on italics or bolding
p.runs[1].text #returns the text of the run

#returns bool if a run contains bold or italic
p.runs[1].bold
p.runs[1].italic

#you can also set bolding/italics to true
p.runs[1].bold = True

#You can also change the values in the word doc
p.runs[1].text = "let's put some new text here"

#you can save after editing by using
d.save('nameOfFile')

#you can also adjust styling
p.style = 'Title'

#create new
d =docx.Document()

d.add_paragraph("String I want to add") #new doc is empty and exists nowhere
p = d.paragraphs[0] #add paragraph to
p.add_run('String to add to previous paragraph') #add run to
d.save('nameOfFile')

#modify existing - you can't directly modify document you have to copy over to a new doc and insert changes
